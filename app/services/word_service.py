import docx
from docx.enum.text import WD_BREAK
from docxcompose.composer import Composer
import os
import shutil
from sqlalchemy.orm import Session
from app.db.models import TemplateDoc, ReportHistory
from app.core.config import TEMPLATES_WORD_DIR, BASE_DIR

def replace_placeholder_robust(p, old_text, new_text):
    """Replaces text in a paragraph while preserving the formatting of the first run."""
    if old_text in p.text:
        updated_text = p.text.replace(old_text, new_text)
        for run in p.runs:
            run.text = ""
        if p.runs:
            p.runs[0].text = updated_text
        else:
            p.add_run(updated_text)

def clean_paragraph_formatting(p):
    """Removes manual font and paragraph overrides to let the style take over."""
    for run in p.runs:
        run.font.size = None
        run.font.bold = None
        run.font.name = None
    
    p.paragraph_format.left_indent = None
    p.paragraph_format.right_indent = None
    p.paragraph_format.first_line_indent = None
    p.paragraph_format.alignment = None

    if p.runs:
        p.runs[0].text = p.runs[0].text.lstrip()

def get_font_size(p):
    """Tries to get the font size of the first non-empty run in a paragraph."""
    for run in p.runs:
        if run.text.strip() and run.font.size:
            return run.font.size.pt
    return None

def get_template_path(db: Session, template_id: int = None):
    if template_id:
        template = db.query(TemplateDoc).filter(TemplateDoc.id == template_id).first()
        if template:
            return os.path.join(TEMPLATES_WORD_DIR, template.filename)
    # Default to Master
    master = db.query(TemplateDoc).filter(TemplateDoc.is_master == True).first()
    if master:
        return os.path.join(TEMPLATES_WORD_DIR, master.filename)
    return os.path.join(BASE_DIR, "template-word.docx") # Fallback

def reformat_docx_smart(template_path, draft_path, output_path, meta):
    """Advanced reformatting logic."""
    if not os.path.exists(template_path):
        raise FileNotFoundError(f"Template not found at {template_path}")
    
    draft = docx.Document(draft_path)
    processed_paragraphs = draft.paragraphs[:]
    
    for p in processed_paragraphs:
        size = get_font_size(p)
        if size:
            if size == 19: p.style = 'Heading 1'
            elif size == 18: p.style = 'Heading 2'
            elif size >= 17: p.style = 'Heading 3'
            elif size == 16: p.style = 'Heading 4'
            elif size == 15: p.style = 'Heading 5'
            elif size == 14: p.style = 'Normal'
            elif size == 13:
                if not p.text.startswith('-'): p.text = f"- {p.text}"
                p.style = 'Normal'
            clean_paragraph_formatting(p)

    temp_draft_path = draft_path + ".cleaned.docx"
    draft.save(temp_draft_path)

    doc = docx.Document(template_path)
    
    # Update Metadata
    for p in doc.paragraphs:
        if "<TIEU DE>" in p.text: replace_placeholder_robust(p, "<TIEU DE>", meta['title'])
        if "<1.0>" in p.text: replace_placeholder_robust(p, "<1.0>", meta['version'])

    for section in doc.sections:
        for p in section.footer.paragraphs:
            if "<Tieu de file>" in p.text: replace_placeholder_robust(p, "<Tieu de file>", meta['title'])
        for p in section.header.paragraphs:
            if "<Tieu de file>" in p.text: replace_placeholder_robust(p, "<Tieu de file>", meta['title'])

    # Clear dummy content from Template (P15+)
    paragraphs_to_remove = doc.paragraphs[15:]
    for p in paragraphs_to_remove:
        p_element = p._element
        p_element.getparent().remove(p_element)

    doc.add_page_break()

    # Use docxcompose to merge
    composer = Composer(doc)
    composer.append(docx.Document(temp_draft_path))
    composer.save(output_path)
    
    if os.path.exists(temp_draft_path):
        os.remove(temp_draft_path)
    return output_path
